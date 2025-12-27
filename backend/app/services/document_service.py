"""
Document Processing Service with Enhanced Excel Intelligence.

This service handles:
- PDF, Word, Excel, Text, CSV file processing
- Intelligent text extraction and chunking
- Excel data analysis and summarization
- Document metadata extraction
"""

import os
import hashlib
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import re

from app.core.config import settings
from app.services.vector_service import vector_service

logger = logging.getLogger(__name__)

# Document processing imports
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger.warning("PyPDF2 not installed - PDF processing disabled")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed - Word processing disabled")

try:
    import openpyxl
    import pandas as pd
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    logger.warning("openpyxl/pandas not installed - Excel processing disabled")

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False


class DocumentProcessor:
    """
    Comprehensive document processing service.
    
    Supports:
    - PDF documents
    - Word documents (.docx, .doc)
    - Excel spreadsheets (.xlsx, .xls)
    - CSV files
    - Plain text files
    """
    
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.ms-excel": "xls",
        "text/csv": "csv",
        "text/plain": "txt",
        "application/octet-stream": "auto"  # Will detect from extension
    }
    
    def __init__(self):
        """Initialize document processor."""
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True, parents=True)
        
        logger.info(f"Document processor initialized. Upload dir: {self.upload_dir}")
        logger.info(f"Available processors - PDF: {HAS_PDF}, DOCX: {HAS_DOCX}, Excel: {HAS_EXCEL}")
    
    def process_document(
        self,
        file_content: bytes,
        filename: str,
        mime_type: str,
        project_id: Optional[int] = None,
        user_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Process an uploaded document.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type of the file
            project_id: Associated project ID
            user_id: Uploading user ID
            
        Returns:
            Dict with processing results and metadata
        """
        start_time = datetime.utcnow()
        
        try:
            # Generate file hash
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Determine file type
            file_ext = Path(filename).suffix.lower()
            doc_type = self._determine_doc_type(mime_type, file_ext)
            
            # Save file temporarily for processing
            temp_path = self._save_temp_file(file_content, filename)
            
            # Extract content based on type
            extraction_result = self._extract_content(temp_path, doc_type, filename)
            
            if not extraction_result.get("success"):
                return extraction_result
            
            # Get extracted text and metadata
            extracted_text = extraction_result.get("text", "")
            doc_metadata = extraction_result.get("metadata", {})
            
            # Chunk the text for embedding
            chunks = self._chunk_text(
                extracted_text,
                filename=filename,
                doc_metadata=doc_metadata
            )
            
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            return {
                "success": True,
                "file_hash": file_hash,
                "filename": filename,
                "doc_type": doc_type,
                "extracted_text": extracted_text,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "word_count": len(extracted_text.split()),
                "char_count": len(extracted_text),
                "metadata": doc_metadata,
                "processing_time": processing_time,
                "is_excel": doc_type in ["xlsx", "xls", "csv"]
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }
    
    def _determine_doc_type(self, mime_type: str, file_ext: str) -> str:
        """Determine document type from MIME type and extension."""
        # First try MIME type
        if mime_type in self.SUPPORTED_TYPES:
            doc_type = self.SUPPORTED_TYPES[mime_type]
            if doc_type != "auto":
                return doc_type
        
        # Fall back to extension
        ext_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".xlsx": "xlsx",
            ".xls": "xls",
            ".csv": "csv",
            ".txt": "txt"
        }
        
        return ext_map.get(file_ext, "txt")
    
    def _save_temp_file(self, content: bytes, filename: str) -> str:
        """Save file content to temporary location."""
        temp_filename = f"temp_{datetime.utcnow().timestamp()}_{filename}"
        temp_path = self.upload_dir / temp_filename
        
        with open(temp_path, "wb") as f:
            f.write(content)
        
        return str(temp_path)
    
    def _extract_content(
        self,
        file_path: str,
        doc_type: str,
        filename: str
    ) -> Dict[str, Any]:
        """Extract content from document based on type."""
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "doc": self._extract_docx,  # Try same as docx
            "xlsx": self._extract_excel,
            "xls": self._extract_excel,
            "csv": self._extract_csv,
            "txt": self._extract_text
        }
        
        extractor = extractors.get(doc_type, self._extract_text)
        return extractor(file_path, filename)
    
    def _extract_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        if not HAS_PDF:
            return {"success": False, "error": "PDF processing not available"}
        
        try:
            text_parts = []
            metadata = {"pages": 0, "filename": filename}
            
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                metadata["pages"] = len(reader.pages)
                
                # Extract PDF metadata
                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title", "")
                    metadata["author"] = reader.metadata.get("/Author", "")
                    metadata["subject"] = reader.metadata.get("/Subject", "")
                
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
            
            return {
                "success": True,
                "text": "\n\n".join(text_parts),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_docx(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract text from Word document."""
        if not HAS_DOCX:
            return {"success": False, "error": "Word processing not available"}
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            metadata = {"filename": filename}
            
            # Extract core properties
            core_props = doc.core_properties
            if core_props:
                metadata["title"] = core_props.title or ""
                metadata["author"] = core_props.author or ""
                metadata["subject"] = core_props.subject or ""
            
            # Extract paragraphs
            current_heading = None
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect headings
                if para.style and para.style.name.startswith("Heading"):
                    current_heading = text
                    text_parts.append(f"\n[Section: {text}]\n")
                else:
                    text_parts.append(text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables, 1):
                text_parts.append(f"\n[Table {table_idx}]")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            metadata["word_count"] = len(" ".join(text_parts).split())
            
            return {
                "success": True,
                "text": "\n".join(text_parts),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_excel(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Extract and analyze Excel spreadsheet data.
        
        Features:
        - Sheet parsing
        - Column recognition
        - Data type detection
        - Basic statistical analysis
        - Trend identification
        """
        if not HAS_EXCEL:
            return {"success": False, "error": "Excel processing not available"}
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            metadata = {
                "filename": filename,
                "sheets": [],
                "total_rows": 0,
                "total_columns": 0,
                "data_summary": {}
            }
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_info = {
                    "name": sheet_name,
                    "rows": sheet.max_row,
                    "columns": sheet.max_column,
                    "headers": [],
                    "data_types": {},
                    "summary_stats": {}
                }
                
                # Extract headers (first row)
                headers = []
                for col in range(1, sheet.max_column + 1):
                    header_value = sheet.cell(row=1, column=col).value
                    header = str(header_value) if header_value else f"Column_{col}"
                    headers.append(header)
                
                sheet_info["headers"] = headers
                
                # Build text representation
                text_parts.append(f"\n[Sheet: {sheet_name}]")
                text_parts.append(f"Columns: {', '.join(headers)}")
                text_parts.append(f"Total Rows: {sheet.max_row - 1} (excluding header)")
                
                # Extract data and analyze
                data_rows = []
                numeric_columns = {col: [] for col in range(len(headers))}
                
                for row_num in range(2, min(sheet.max_row + 1, 1002)):  # Limit to 1000 rows
                    row_data = []
                    for col_num, header in enumerate(headers, 1):
                        cell_value = sheet.cell(row=row_num, column=col_num).value
                        
                        # Track numeric values for analysis
                        if isinstance(cell_value, (int, float)):
                            numeric_columns[col_num - 1].append(cell_value)
                        
                        row_data.append(str(cell_value) if cell_value is not None else "")
                    
                    if any(row_data):
                        data_rows.append(row_data)
                
                # Add sample data to text
                text_parts.append("\nSample Data (first 20 rows):")
                for i, row in enumerate(data_rows[:20]):
                    text_parts.append(" | ".join(row))
                
                # Calculate summary statistics for numeric columns
                for col_idx, values in numeric_columns.items():
                    if len(values) > 0:
                        col_name = headers[col_idx]
                        stats = self._calculate_column_stats(values)
                        sheet_info["summary_stats"][col_name] = stats
                        
                        text_parts.append(f"\n[Statistics for '{col_name}':]")
                        text_parts.append(f"  Count: {stats['count']}")
                        text_parts.append(f"  Sum: {stats['sum']:.2f}")
                        text_parts.append(f"  Average: {stats['mean']:.2f}")
                        text_parts.append(f"  Min: {stats['min']:.2f}")
                        text_parts.append(f"  Max: {stats['max']:.2f}")
                        
                        # Detect anomalies
                        if stats.get("anomalies"):
                            text_parts.append(f"  Potential Anomalies: {len(stats['anomalies'])} detected")
                
                metadata["sheets"].append(sheet_info)
                metadata["total_rows"] += sheet.max_row
                metadata["total_columns"] = max(metadata["total_columns"], sheet.max_column)
            
            return {
                "success": True,
                "text": "\n".join(text_parts),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _calculate_column_stats(self, values: List[float]) -> Dict[str, Any]:
        """Calculate statistics for a numeric column."""
        if not values:
            return {}
        
        count = len(values)
        total = sum(values)
        mean = total / count
        min_val = min(values)
        max_val = max(values)
        
        # Calculate standard deviation
        variance = sum((x - mean) ** 2 for x in values) / count
        std_dev = variance ** 0.5
        
        # Detect anomalies (values > 2 std from mean)
        anomalies = []
        if std_dev > 0:
            for i, v in enumerate(values):
                z_score = abs(v - mean) / std_dev
                if z_score > 2:
                    anomalies.append({"row": i + 2, "value": v, "z_score": z_score})
        
        return {
            "count": count,
            "sum": total,
            "mean": mean,
            "min": min_val,
            "max": max_val,
            "std_dev": std_dev,
            "anomalies": anomalies[:10]  # Limit to 10 anomalies
        }
    
    def _extract_csv(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract data from CSV file."""
        try:
            # Detect encoding
            encoding = "utf-8"
            if HAS_CHARDET:
                with open(file_path, "rb") as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result.get("encoding", "utf-8")
            
            df = pd.read_csv(file_path, encoding=encoding)
            
            text_parts = []
            metadata = {
                "filename": filename,
                "rows": len(df),
                "columns": list(df.columns),
                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            
            text_parts.append(f"[CSV File: {filename}]")
            text_parts.append(f"Columns: {', '.join(df.columns)}")
            text_parts.append(f"Total Rows: {len(df)}")
            
            # Add sample data
            text_parts.append("\nSample Data (first 20 rows):")
            text_parts.append(df.head(20).to_string())
            
            # Add statistics for numeric columns
            numeric_df = df.select_dtypes(include=['number'])
            if len(numeric_df.columns) > 0:
                text_parts.append("\n[Numeric Column Statistics:]")
                text_parts.append(numeric_df.describe().to_string())
            
            return {
                "success": True,
                "text": "\n".join(text_parts),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Extract text from plain text file."""
        try:
            # Detect encoding
            encoding = "utf-8"
            if HAS_CHARDET:
                with open(file_path, "rb") as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result.get("encoding", "utf-8") or "utf-8"
            
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                content = f.read()
            
            return {
                "success": True,
                "text": content,
                "metadata": {
                    "filename": filename,
                    "encoding": encoding,
                    "char_count": len(content),
                    "word_count": len(content.split()),
                    "line_count": content.count("\n") + 1
                }
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _chunk_text(
        self,
        text: str,
        filename: str = "",
        doc_metadata: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks for embedding.
        
        Args:
            text: Full document text
            filename: Source filename
            doc_metadata: Document metadata
            
        Returns:
            List of chunk dictionaries
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        words = text.split()
        chunk_size = settings.CHUNK_SIZE
        overlap = settings.CHUNK_OVERLAP
        
        if len(words) <= chunk_size:
            # Document fits in single chunk
            return [{
                "content": text,
                "chunk_index": 0,
                "token_count": len(words),
                "start_word": 0,
                "end_word": len(words),
                "filename": filename,
                "page_number": doc_metadata.get("pages", 1) if doc_metadata else None
            }]
        
        # Create overlapping chunks
        i = 0
        chunk_index = 0
        
        while i < len(words):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            
            # Try to detect page/section from content
            page_number = self._detect_page_number(chunk_text)
            section_title = self._detect_section_title(chunk_text)
            
            chunks.append({
                "content": chunk_text,
                "chunk_index": chunk_index,
                "token_count": len(chunk_words),
                "start_word": i,
                "end_word": min(i + chunk_size, len(words)),
                "filename": filename,
                "page_number": page_number,
                "section_title": section_title
            })
            
            i += chunk_size - overlap
            chunk_index += 1
        
        return chunks
    
    def _detect_page_number(self, text: str) -> Optional[int]:
        """Detect page number from chunk text."""
        patterns = [
            r'\[Page\s+(\d+)\]',
            r'Page\s+(\d+)\s+of',
            r'- (\d+) -'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        
        return None
    
    def _detect_section_title(self, text: str) -> Optional[str]:
        """Detect section title from chunk text."""
        patterns = [
            r'\[Section:\s*([^\]]+)\]',
            r'\[Sheet:\s*([^\]]+)\]',
            r'^#\s+(.+)$'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def save_document_file(
        self,
        file_content: bytes,
        filename: str,
        document_id: int
    ) -> str:
        """
        Save document file to permanent storage.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            document_id: Database document ID
            
        Returns:
            Path to saved file
        """
        # Create document-specific directory
        doc_dir = self.upload_dir / str(document_id)
        doc_dir.mkdir(exist_ok=True, parents=True)
        
        # Save file
        file_path = doc_dir / filename
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        return str(file_path)
    
    def delete_document_file(self, document_id: int) -> bool:
        """Delete document files from storage."""
        try:
            doc_dir = self.upload_dir / str(document_id)
            if doc_dir.exists():
                import shutil
                shutil.rmtree(doc_dir)
            return True
        except Exception as e:
            logger.error(f"Failed to delete document files: {e}")
            return False


# Global instance
document_processor = DocumentProcessor()
